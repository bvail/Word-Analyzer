#!/usr/bin/env python3

import docx
from string import punctuation
from unidecode import unidecode

class WordAnalyzer():
    
    def __init__(self, doc_name):
        self._doc = docx.Document(doc_name)
        self._word_list = []
    

    def analyze_doc(self):
        
        for i in range(len(self._doc.paragraphs)):
        
            string = self._doc.paragraphs[i].text
            string = unidecode(string)
            _clean_string = self.remove_punctuation(string)
            _lowercase_string = self.remove_capitalization(_clean_string)
            _string_list = self.split_string(_lowercase_string)
            self.create_word_list(_string_list)
            self.sort_word_list()
        
        
        
    
    # removes punctuation from string
    def remove_punctuation(self, string):
        _list = string.split('-') # split string by dashes
        string = ' '.join(_list)  # separates dashes so compound words are split into 2 words
        return ''.join(c for c in string if c not in punctuation)
    
    # removes capitalization from string
    def remove_capitalization(self, string):
        return string.lower()
    
    
    # splits string into individual words in list
    def split_string(self, string):
        return string.split()
    
    
    # creates 2d list of words from the document
    # in row 1, each word is added only once
    # in row 2, the frequency of each word is tallied
    def create_word_list(self, string_list):
        
       
        for i in range(len(string_list)):
            
            update = False
            
            # if the word in string_list is already in word_list, updates tally
            for j in range(len(self._word_list)):
                if string_list[i] == self._word_list[j][0]:
                    self._word_list[j][1] += 1
                    update = True
            
            # if the word in string_list is not in word_list, adds to word_list and start tally at 1
            if update == False:
                self._word_list.insert(0, [string_list[i], 1])
    
    
    def sort_word_list(self):
        self._word_list.sort(key = sort_second, reverse = True)
   
   
    def print_word_list(self):
        for i in range(len(self._word_list)):
            print(self._word_list[i][0], '-', self._word_list[i][1])
    
    def get_word_list(self):
        return self._word_list
            
    
    def word_count(self):
        count = 0
        for i in range(len(self._word_list)):
            count += self._word_list[i][1]
        return count
            
   
    def unique_word_count(self):
        return len(self._word_list)
    
    # Search _word_list for key word, returns count
    def search_word_list(self, key):
        
        key = unidecode(key)
        key = self.remove_capitalization(key)
        
        for i in range(len(self._word_list)):
            if key == self._word_list[i][0]:
                return self._word_list[i][1]
            else:
                return None
   
 
 # function to sort 2d array by ""count row    
def sort_second(elem):
    return elem[1]
            
            
    
        
    
    


if __name__ == '__main__': main()